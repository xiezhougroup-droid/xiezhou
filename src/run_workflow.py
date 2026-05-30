#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

CODEX_PYTHON_PACKAGES = Path("/Users/houzhou/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/lib/python3.12/site-packages")
if CODEX_PYTHON_PACKAGES.exists():
    sys.path.insert(0, str(CODEX_PYTHON_PACKAGES))

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


TOOL_DIR = Path(__file__).resolve().parents[1]
PROJECT_DIR = TOOL_DIR.parent
CONFIG_DIR = TOOL_DIR / "config"

BOQ_HEADERS = [
    "原始行号",
    "原始编号",
    "原始材料名称",
    "标准材料名称",
    "规格型号",
    "单位",
    "数量",
    "参考单价",
    "参考合价",
    "原始类别",
    "询价状态",
    "材料大类",
    "建议供应商包",
    "需人工确认",
    "确认原因",
]

QUOTE_HEADERS = [
    "序号",
    "名称",
    "规格型号",
    "单位",
    "数量",
    "报价单价",
    "报价合价",
    "备注",
    "追溯号",
    "原始材料名称",
]

INDEX_HEADERS = [
    "文件名",
    "相对路径",
    "所属专业目录",
    "图纸标题或标题摘要",
    "页码或图号",
    "关键词",
    "建议匹配供应商包",
    "是否需人工确认",
    "确认原因",
    "建议挂接状态",
    "文件大小KB",
    "识别方式",
]

ATTACH_HEADERS = [
    "文件名",
    "相对路径",
    "所属专业目录",
    "图纸标题或标题摘要",
    "页码或图号",
    "关键词",
    "建议匹配供应商包",
    "建议挂接状态",
    "是否需人工确认",
    "确认原因",
    "建议用途",
]

MANUAL_DRAWING_HEADERS = [
    "文件名",
    "相对路径",
    "所属专业目录",
    "图纸标题或标题摘要",
    "页码或图号",
    "关键词",
    "建议匹配供应商包",
    "是否需人工确认",
    "确认原因",
    "建议挂接状态",
]


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def clean(value: Any) -> str:
    return "" if value is None else str(value).strip()


def resolve_path(base_dir: Path, configured: str) -> Path:
    path = Path(configured)
    return path if path.is_absolute() else base_dir / path


def read_public_info(path: Path) -> dict[str, str]:
    if not path.exists():
        return {}
    if path.suffix.lower() in [".xlsx", ".xlsm"]:
        wb = load_workbook(path, data_only=True)
        values: dict[str, str] = {}
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cleaned = [clean(cell) for cell in row]
                non_empty = [cell for cell in cleaned if cell]
                if len(non_empty) >= 2:
                    values[non_empty[0].rstrip("：:")] = non_empty[1]
        return values
    if path.suffix.lower() == ".rtf":
        try:
            result = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", str(path)],
                check=True,
                capture_output=True,
                text=True,
            )
            text = result.stdout
        except Exception:
            text = path.read_text(encoding="utf-8", errors="ignore")
    else:
        text = path.read_text(encoding="utf-8", errors="ignore")

    values: dict[str, str] = {}
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if "：" in line:
            key, value = line.split("：", 1)
        elif ":" in line:
            key, value = line.split(":", 1)
        else:
            continue
        values[key.strip()] = value.strip()
    return values


def row_to_dict(headers: list[str], row: tuple[Any, ...]) -> dict[str, Any]:
    return {headers[i]: row[i] if i < len(row) else None for i in range(len(headers))}


def trace_id(row: dict[str, Any]) -> str:
    original_row = clean(row.get("原始行号"))
    original_code = clean(row.get("原始编号"))
    return f"{original_row}｜{original_code}" if original_code else original_row


def read_boq_items(boq_path: Path, sheet_name: str, source_supplier_package: str) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    wb = load_workbook(boq_path, data_only=True)
    ws = wb[sheet_name]
    headers = [cell.value for cell in ws[1]]
    if headers[: len(BOQ_HEADERS)] != BOQ_HEADERS:
        raise RuntimeError("底表字段与预期不一致，请检查正式询价清单底表。")

    quote_items: list[dict[str, Any]] = []
    manual_items: list[dict[str, Any]] = []
    for values in ws.iter_rows(min_row=2, values_only=True):
        item = row_to_dict(BOQ_HEADERS, values)
        if clean(item.get("询价状态")) != "拟询价":
            continue
        if clean(item.get("建议供应商包")) != source_supplier_package:
            continue
        if clean(item.get("需人工确认")) == "是":
            manual_items.append(item)
        else:
            quote_items.append(item)
    return quote_items, manual_items


def title_from_filename(path: Path) -> str:
    stem = path.stem
    if "@" in stem:
        return stem.split("@", 1)[1].strip()
    if "_" in stem:
        return stem.split("_", 1)[1].strip()
    if " " in stem and re.search(r"(建施|暖施|电施|水施|结施|装配施)", stem):
        return stem.split(" ", 1)[1].strip()
    parts = stem.split("-")
    if len(parts) >= 3:
        return "-".join(parts[2:]).strip()
    return stem


def drawing_no(path: Path) -> str:
    stem = path.stem
    patterns = [
        r"(建施[-A-Z]*\d+[a-zA-Z]?)",
        r"(暖施[-A-Z]*\d+[a-zA-Z]?)",
        r"(电施[-A-Z]*\d+[a-zA-Z]?)",
        r"(水施[-A-Z]*\d+[a-zA-Z]?)",
        r"(结施[A-Z]\d+)",
        r"(装配施\d+)",
        r"(总施\d+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, stem)
        if match:
            return match.group(1)
    return ""


def specialty_from_path(path: Path) -> str:
    for part in path.parts:
        if part.startswith("施工图设计文件-"):
            return part.replace("施工图设计文件-", "")
    return path.parent.name


def collect_keywords(text: str, package_cfg: dict[str, Any]) -> str:
    candidates = list(dict.fromkeys(
        package_cfg.get("direct_attachment_keywords", [])
        + package_cfg.get("candidate_attachment_keywords", [])
        + ["门窗", "铝合金", "节点详图", "总说明", "构造做法", "平面图", "立面图", "材料表", "图纸目录"]
    ))
    return "、".join(keyword for keyword in candidates if keyword in text)


def package_hint(specialty: str) -> str:
    if "建筑" in specialty:
        return "建筑/装饰相关供应商包"
    if "暖通" in specialty:
        return "暖通空调/通风设备相关供应商包"
    if "电气" in specialty:
        return "电气设备材料供应商包"
    if "给排水" in specialty:
        return "给排水/管材管件相关供应商包"
    if "结构" in specialty:
        return "结构/土建相关供应商包"
    if "装配式" in specialty:
        return "装配式相关供应商包"
    return ""


def classify_drawing(path: Path, specialty: str, title: str, keywords: str, package_cfg: dict[str, Any]) -> tuple[str, str, str, str]:
    source = f"{path.name} {title} {keywords}"
    direct_keywords = package_cfg.get("direct_attachment_keywords", [])
    candidate_keywords = package_cfg.get("candidate_attachment_keywords", [])
    direct_specialties = package_cfg.get("direct_attachment_specialties", [])
    candidate_specialties = package_cfg.get("candidate_attachment_specialties", [])
    source_package = package_cfg["source_supplier_package"]

    if any(keyword in source for keyword in direct_keywords) and any(sp in specialty for sp in direct_specialties):
        return source_package, "否", "文件名明确命中正式附件关键词", "建议挂接"

    if any(keyword in source for keyword in candidate_keywords) and any(sp in specialty for sp in candidate_specialties):
        return source_package, "是", "可能关联本采购包，但需确认是否纳入正式附件", "候选，需确认"

    if "建筑" in specialty and any(keyword in source for keyword in ["节点详图", "构造做法", "总说明", "立面图"]):
        return package_hint(specialty), "是", "可能包含相关技术要求，但文件名无法明确判断", "不挂接，保留人工复核"

    return package_hint(specialty), "否", "", "不挂接"


def index_drawings(drawings_dir: Path, base_dir: Path, package_cfg: dict[str, Any]) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    records: list[dict[str, Any]] = []
    for path in sorted(drawings_dir.rglob("*.pdf")):
        rel = path.relative_to(base_dir)
        specialty = specialty_from_path(path.relative_to(drawings_dir))
        title = title_from_filename(path)
        keywords = collect_keywords(f"{path.name} {title}", package_cfg)
        suggested_package, needs_manual, reason, attach_status = classify_drawing(path, specialty, title, keywords, package_cfg)
        records.append({
            "文件名": path.name,
            "相对路径": str(rel),
            "所属专业目录": specialty,
            "图纸标题或标题摘要": title,
            "页码或图号": drawing_no(path),
            "关键词": keywords,
            "建议匹配供应商包": suggested_package,
            "是否需人工确认": needs_manual,
            "确认原因": reason,
            "建议挂接状态": attach_status,
            "文件大小KB": round(path.stat().st_size / 1024, 1),
            "识别方式": "文件名规则识别",
        })

    source_package = package_cfg["source_supplier_package"]
    attachments = [
        {**record, "建议用途": "作为报价邀请包正式附件" if record["建议挂接状态"] == "建议挂接" else "候选附件，需人工确认"}
        for record in records
        if record["建议匹配供应商包"] == source_package and record["建议挂接状态"] in ["建议挂接", "候选，需确认"]
    ]
    manual = [record for record in records if record["是否需人工确认"] == "是"]
    return records, attachments, manual


def write_sheet(ws, headers: list[str], rows: list[list[Any]], title: str) -> None:
    dark = PatternFill("solid", fgColor="1F4E78")
    blue = PatternFill("solid", fgColor="5B9BD5")
    thin = Side(style="thin", color="D9E2F3")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    widths = {
        "文件名": 46,
        "相对路径": 72,
        "所属专业目录": 18,
        "图纸标题或标题摘要": 48,
        "页码或图号": 18,
        "关键词": 32,
        "建议匹配供应商包": 30,
        "是否需人工确认": 16,
        "确认原因": 52,
        "建议挂接状态": 18,
        "建议用途": 44,
        "文件大小KB": 14,
        "识别方式": 18,
        "项目": 28,
        "内容": 72,
    }

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    ws.cell(1, 1, title)
    ws.cell(1, 1).font = Font(bold=True, size=14, color="FFFFFF")
    ws.cell(1, 1).fill = dark
    ws.cell(1, 1).alignment = Alignment(horizontal="center")
    header_row = 3

    for col, header in enumerate(headers, 1):
        cell = ws.cell(header_row, col, header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = blue
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

    for r_idx, row in enumerate(rows, header_row + 1):
        for c_idx, value in enumerate(row, 1):
            cell = ws.cell(r_idx, c_idx, value)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border

    ws.freeze_panes = ws.cell(header_row + 1, 1)
    if rows:
        ws.auto_filter.ref = f"A{header_row}:{get_column_letter(len(headers))}{header_row + len(rows)}"
    ws.sheet_view.showGridLines = False
    for idx, header in enumerate(headers, 1):
        ws.column_dimensions[get_column_letter(idx)].width = widths.get(header, 18)


def write_index_workbook(output_path: Path, records: list[dict[str, Any]], attachments: list[dict[str, Any]], manual: list[dict[str, Any]], package_cfg: dict[str, Any]) -> None:
    wb = Workbook()
    ws_summary = wb.active
    ws_summary.title = "处理说明"
    specialty_counter = Counter(record["所属专业目录"] for record in records)
    summary_rows = [
        ["供应商包", package_cfg["display_package_name"]],
        ["来源供应商包字段", package_cfg["source_supplier_package"]],
        ["PDF总数", len(records)],
        ["图纸索引方式", "文件名规则识别"],
        ["附件匹配数量", len(attachments)],
        ["建议直接挂接", sum(1 for item in attachments if item["建议挂接状态"] == "建议挂接")],
        ["候选需确认", sum(1 for item in attachments if item["建议挂接状态"] == "候选，需确认")],
        ["图纸人工确认项", len(manual)],
        ["说明", "本工具不修改原始图纸，不覆盖原始文件。"],
    ]
    for specialty, count in specialty_counter.most_common():
        summary_rows.append([f"{specialty} PDF数量", count])
    write_sheet(ws_summary, ["项目", "内容"], summary_rows, "处理说明")

    ws_index = wb.create_sheet("图纸资料索引")
    write_sheet(ws_index, INDEX_HEADERS, [[record[h] for h in INDEX_HEADERS] for record in records], "图纸资料索引表")

    ws_attach = wb.create_sheet("附件匹配表")
    write_sheet(ws_attach, ATTACH_HEADERS, [[record[h] for h in ATTACH_HEADERS] for record in attachments], "供应商包附件匹配表")

    ws_manual = wb.create_sheet("图纸需人工确认项")
    write_sheet(ws_manual, MANUAL_DRAWING_HEADERS, [[record[h] for h in MANUAL_DRAWING_HEADERS] for record in manual], "图纸匹配需人工确认项")

    wb.save(output_path)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="BFBFBF") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, bold=False, size=11, color=None, font="SimSun") -> None:
    run.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if font == "SimHei" else "宋体")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    style_run(r, bold=True, size=12 if level == 1 else 11, color=(31, 78, 121), font="SimHei")


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    style_run(r, size=10.5)


def add_label_para(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(label)
    style_run(r, bold=True, size=11)
    r2 = p.add_run(value)
    style_run(r2, size=11)


def format_docx_table(table, widths_cm: list[float], header_fill="E7EEF7", font_size=8) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, header_fill)
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    style_run(run, bold=(r_idx == 0), size=font_size)


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(6)


def write_quote_docx(output_path: Path, quote_items: list[dict[str, Any]], manual_items: list[dict[str, Any]], attachments: list[dict[str, Any]], public_info: dict[str, str], package_cfg: dict[str, Any]) -> None:
    direct_attachments = [item for item in attachments if item["建议挂接状态"] == "建议挂接"]
    candidate_attachments = [item for item in attachments if item["建议挂接状态"] == "候选，需确认"]
    display_name = package_cfg["display_package_name"]

    doc = Document()
    setup_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("报价邀请函")
    style_run(run, bold=True, size=18, font="SimHei")

    p = doc.add_paragraph()
    r = p.add_run("致：各受邀报价单位")
    style_run(r, bold=True, size=11)
    add_body(doc, package_cfg["intro_scope"])

    add_heading(doc, "一、项目简介")
    add_label_para(doc, "项目名称：", public_info.get("项目名称") or "待补充")
    add_label_para(doc, "工程名称：", display_name)
    add_label_para(doc, "工程地点：", public_info.get("工程地点") or "待补充")
    add_label_para(doc, "报价范围：", package_cfg["quote_scope"])

    add_heading(doc, "二、范围及承包方式")
    add_body(doc, public_info.get("承包方式默认口径") or "本工程暂按包工包料、综合单价报价方式考虑。报价应包含完成本报价范围所需的材料、加工制作、运输、装卸、成品保护、配合验收、税费、管理费、利润及完成本项工作的其他必要费用。")
    add_body(doc, "本文件已挂接匹配到的正式图纸附件，报价单位应结合已挂接图纸和报价清单进行核对。尚未完全覆盖的现场条件、深化设计、品牌档次、检测要求、样品确认、施工配合及其他专项要求，后续由采购/项目团队补充确认。")

    add_heading(doc, "三、质量要求")
    add_body(doc, public_info.get("质量要求默认口径") or package_cfg["quality_requirement"])

    add_heading(doc, "四、工期要求")
    add_body(doc, public_info.get("工期要求默认口径") or "工期及供货节点须满足项目总体进度安排和甲方书面通知要求。报价单位应在报价文件中明确加工周期、供货周期、安装或配合周期，以及影响工期的前置条件。")

    add_heading(doc, "五、报价文件需包含内容")
    quote_file_content = public_info.get("报价文件需包含内容默认口径")
    if quote_file_content:
        add_body(doc, quote_file_content)
    else:
        for text in [
            "公司简介、营业执照、资质或类似能力证明文件；",
            "近三年类似工程业绩或供货业绩；",
            "按本邀请函附件格式填写的报价清单，报价单价和合价应完整；",
            "主要材料、产品说明、品牌档次及技术响应说明；",
            "联系人、联系电话、报价有效期、供货周期及需采购方确认的偏离事项。",
        ]:
            add_body(doc, text)

    add_heading(doc, "六、付款方式")
    add_body(doc, public_info.get("付款方式默认口径") or "付款方式暂按公司标准专业分包/采购付款条款执行，具体付款节点、比例、发票类型、质保金及结算方式以后续合同谈判及正式合同约定为准。报价单位可在报价文件中说明建议付款条件。")

    add_heading(doc, "七、联系人")
    add_label_para(doc, "联系人：", public_info.get("联系人") or "待补充")
    add_label_para(doc, "联系电话：", public_info.get("联系电话") or "待补充")
    add_label_para(doc, "报价截止时间：", "待补充")
    add_label_para(doc, "报价文件提交方式：", "待补充")

    add_heading(doc, "八、附件说明")
    add_body(doc, f"附件一：{display_name}报价清单")
    add_body(doc, "附件二：已匹配并挂接的图纸资料，详见本文件第十条正式附件清单。")
    add_body(doc, "附件三：待人工确认附件及需人工确认项提示表")

    add_heading(doc, "九、技术要求 / 图纸说明")
    add_body(doc, f"本次已根据图纸资料索引匹配并挂接与{display_name}相关的正式图纸。报价单位应结合下表图纸、报价清单及后续补充技术文件进行报价，{package_cfg['focus_check']}")
    tech_table = doc.add_table(rows=1, cols=5)
    for i, header in enumerate(["序号", "图纸名称", "图号/页码", "标题摘要", "建议供应商重点核对内容"]):
        tech_table.rows[0].cells[i].text = header
    for idx, item in enumerate(direct_attachments, 1):
        values = [str(idx), item["文件名"], item["页码或图号"], item["图纸标题或标题摘要"], package_cfg["focus_check"]]
        cells = tech_table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value or ""
    format_docx_table(tech_table, [1.0, 5.3, 2.1, 4.0, 6.5], font_size=8)

    add_heading(doc, "十、正式附件清单")
    for idx, item in enumerate(direct_attachments, 1):
        add_body(doc, f"附件二-{idx}：{item['文件名']}（{item['页码或图号'] or item['图纸标题或标题摘要']}）")

    if candidate_attachments:
        add_heading(doc, "十一、待人工确认附件")
        add_body(doc, "以下资料可能涉及本报价范围，暂不作为正式附件挂接；需采购、技术或项目团队确认是否纳入本次报价范围。")
        candidate_table = doc.add_table(rows=1, cols=5)
        for i, header in enumerate(["序号", "候选图纸名称", "图号/页码", "标题摘要", "待确认原因"]):
            candidate_table.rows[0].cells[i].text = header
        for idx, item in enumerate(candidate_attachments, 1):
            values = [str(idx), item["文件名"], item["页码或图号"], item["图纸标题或标题摘要"], item["确认原因"]]
            cells = candidate_table.add_row().cells
            for i, value in enumerate(values):
                cells[i].text = value or ""
        format_docx_table(candidate_table, [1.0, 5.3, 2.1, 4.0, 6.5], header_fill="FCE4D6", font_size=8)

    quote_section = doc.add_section(WD_SECTION.NEW_PAGE)
    quote_section.orientation = WD_ORIENT.LANDSCAPE
    quote_section.page_width, quote_section.page_height = quote_section.page_height, quote_section.page_width
    quote_section.top_margin = Cm(1.5)
    quote_section.bottom_margin = Cm(1.5)
    quote_section.left_margin = Cm(1.3)
    quote_section.right_margin = Cm(1.3)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = heading.add_run(f"附件一：{display_name}报价清单")
    style_run(hr, bold=True, size=14, font="SimHei")

    quote_table = doc.add_table(rows=1, cols=len(QUOTE_HEADERS))
    for i, header in enumerate(QUOTE_HEADERS):
        quote_table.rows[0].cells[i].text = header
    for idx, item in enumerate(quote_items, 1):
        values = [
            str(idx),
            clean(item.get("标准材料名称")),
            clean(item.get("规格型号")),
            clean(item.get("单位")),
            clean(item.get("数量")),
            "",
            "",
            "",
            trace_id(item),
            clean(item.get("原始材料名称")),
        ]
        cells = quote_table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    format_docx_table(quote_table, [1.0, 3.5, 3.0, 1.0, 1.6, 1.8, 1.8, 2.2, 2.5, 4.2], font_size=8)
    add_body(doc, "说明：报价单价、报价合价及备注由报价单位填写；追溯号用于后续报价回填，请勿删除或改写。")

    if manual_items:
        add_heading(doc, "附件三：需人工确认项提示表")
        add_body(doc, package_cfg["manual_item_intro"])
        manual_headers = ["原始行号", "名称", "规格型号", "单位", "数量", "确认原因", "追溯号"]
        manual_table = doc.add_table(rows=1, cols=len(manual_headers))
        for i, header in enumerate(manual_headers):
            manual_table.rows[0].cells[i].text = header
        for item in manual_items:
            values = [
                clean(item.get("原始行号")),
                clean(item.get("标准材料名称")),
                clean(item.get("规格型号")) or "待补充",
                clean(item.get("单位")),
                clean(item.get("数量")),
                clean(item.get("确认原因")),
                trace_id(item),
            ]
            cells = manual_table.add_row().cells
            for i, value in enumerate(values):
                cells[i].text = value
        format_docx_table(manual_table, [1.5, 4.2, 3.2, 1.2, 1.8, 5.2, 3.2], header_fill="FCE4D6", font_size=8)

    doc.save(output_path)


def write_log(path: Path, lines: list[str]) -> None:
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="采购询价本地工作流")
    parser.add_argument("--package", default=None, help="packages.json 中的供应商包 key")
    parser.add_argument("--base-dir", default=str(PROJECT_DIR), help="项目根目录")
    parser.add_argument("--workflow-config", default=str(CONFIG_DIR / "workflow.json"), help="workflow.json 路径")
    parser.add_argument("--packages-config", default=str(CONFIG_DIR / "packages.json"), help="packages.json 路径")
    parser.add_argument("--boq", default=None, help="临时指定正式询价清单底表路径")
    parser.add_argument("--public-info", default=None, help="临时指定项目公共信息表路径")
    parser.add_argument("--drawings", default=None, help="临时指定图纸文件夹路径")
    parser.add_argument("--output-dir", default=None, help="临时指定输出目录")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    base_dir = Path(args.base_dir).resolve()
    workflow_cfg = load_json(Path(args.workflow_config))
    packages_cfg = load_json(Path(args.packages_config))
    package_key = args.package or packages_cfg["default_package_key"]
    package_cfg = packages_cfg["packages"][package_key]

    boq_path = resolve_path(base_dir, args.boq or workflow_cfg["boq_file"])
    public_info_path = resolve_path(base_dir, args.public_info or workflow_cfg["public_info_file"])
    drawings_dir = resolve_path(base_dir, args.drawings or workflow_cfg["drawings_dir"])
    output_root = resolve_path(base_dir, args.output_dir or workflow_cfg["output_dir"])
    run_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = output_root / f"{run_id}_{package_key}"
    output_dir.mkdir(parents=True, exist_ok=True)

    public_info = read_public_info(public_info_path)
    quote_items, manual_items = read_boq_items(boq_path, workflow_cfg["boq_sheet"], package_cfg["source_supplier_package"])
    drawing_records, attachments, drawing_manual = index_drawings(drawings_dir, base_dir, package_cfg)

    index_output = output_dir / "图纸资料索引及附件匹配表.xlsx"
    write_index_workbook(index_output, drawing_records, attachments, drawing_manual, package_cfg)

    docx_output = output_dir / f"{package_cfg['display_package_name']}.docx"
    write_quote_docx(docx_output, quote_items, manual_items, attachments, public_info, package_cfg)

    direct_count = sum(1 for item in attachments if item["建议挂接状态"] == "建议挂接")
    candidate_count = sum(1 for item in attachments if item["建议挂接状态"] == "候选，需确认")
    log_lines = [
        f"运行时间：{run_id}",
        f"供应商包key：{package_key}",
        f"输出目录：{output_dir}",
        f"底表：{boq_path}",
        f"项目公共信息：{public_info_path}",
        f"图纸目录：{drawings_dir}",
        f"正式报价项：{len(quote_items)}",
        f"清单需人工确认项：{len(manual_items)}",
        f"图纸PDF索引：{len(drawing_records)}",
        f"附件匹配总数：{len(attachments)}",
        f"建议直接挂接附件：{direct_count}",
        f"候选需确认附件：{candidate_count}",
        f"图纸人工确认项：{len(drawing_manual)}",
        f"报价邀请包：{docx_output}",
        f"图纸索引表：{index_output}",
    ]
    write_log(output_dir / "处理日志.txt", log_lines)

    print("\n".join(log_lines))


if __name__ == "__main__":
    main()
