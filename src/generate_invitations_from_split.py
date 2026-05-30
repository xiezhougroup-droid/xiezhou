#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

CODEX_PYTHON_PACKAGES = Path("/Users/houzhou/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/lib/python3.12/site-packages")
if CODEX_PYTHON_PACKAGES.exists():
    sys.path.insert(0, str(CODEX_PYTHON_PACKAGES))

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


TOOL_DIR = Path(__file__).resolve().parents[1]
PROJECT_DIR = TOOL_DIR.parent
CONFIG_DIR = TOOL_DIR / "config"

SOURCE_SHEET = "采购包拆分总表"
HEADER_ROW = 3

QUOTE_HEADERS = [
    "序号",
    "来源行号",
    "追溯号",
    "名称",
    "规格型号",
    "单位",
    "数量",
    "品牌/厂家",
    "税率",
    "供货周期",
    "付款条件",
    "报价单价",
    "报价合价",
    "报价备注",
]

CONFIRM_HEADERS = [
    "序号",
    "来源行号",
    "追溯号",
    "名称",
    "规格型号",
    "单位",
    "数量",
    "是否需人工确认",
    "是否可直接发供应商询价",
    "询价前需补充信息",
    "备注",
]


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def clean(value: Any) -> str:
    return "" if value is None else str(value).strip()


def resolve_path(base_dir: Path, value: str) -> Path:
    path = Path(value)
    return path if path.is_absolute() else base_dir / path


def read_rows(path: Path) -> list[dict[str, Any]]:
    wb = load_workbook(path, data_only=True, read_only=True)
    ws = wb[SOURCE_SHEET]
    headers = [cell.value for cell in ws[HEADER_ROW]]
    rows: list[dict[str, Any]] = []
    for values in ws.iter_rows(min_row=HEADER_ROW + 1, values_only=True):
        row = dict(zip(headers, values))
        if clean(row.get("询价采购包")):
            rows.append(row)
    return rows


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="BFBFBF") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, bold=False, size=10.5, color=None, font="SimSun") -> None:
    run.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if font == "SimHei" else "宋体")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, bold=True, size=12, color=(31, 78, 121), font="SimHei")


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, size=10.5)


def add_label(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label)
    style_run(r, bold=True)
    r2 = p.add_run(value or "待补充")
    style_run(r2)


def format_docx_table(table, header_fill="E7EEF7", font_size=8) -> None:
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    style_run(run, bold=(row_idx == 0), size=font_size)


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.1


def ready_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [row for row in rows if clean(row.get("是否可直接发供应商询价")) == "是"]


def pending_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [row for row in rows if clean(row.get("是否可直接发供应商询价")) != "是"]


def generic_template(package: str, rows: list[dict[str, Any]]) -> dict[str, str]:
    supplier_type = Counter(clean(row.get("供应商类型")) for row in rows).most_common(1)[0][0]
    category = Counter(clean(row.get("询价大类")) for row in rows).most_common(1)[0][0]
    title_suffix = "报价邀请包" if ready_rows(rows) else "报价准备包"
    if "设备" in supplier_type or "配电" in package or "灯具" in package:
        contract_mode = "设备/材料供应，是否包含深化、安装调试配合及验收资料待确认。"
    elif any(word in package for word in ["门窗", "栏杆", "风管", "脚手架"]):
        contract_mode = "材料供应或材料供应+安装/加工配合待确认。"
    else:
        contract_mode = "材料供应，暂按综合到场单价报价。"
    return {
        "package_title": f"{package}{title_suffix}",
        "contract_mode": contract_mode,
        "scope": f"本次报价范围为《采购包拆分总表》中归入“{package}”的{category}相关材料/设备，具体名称、规格、单位及数量详见报价清单。",
        "quality": "质量标准须满足国家、地方及行业现行规范、设计文件、建设单位及总包单位相关要求，并满足项目验收和使用要求。报价单位应明确品牌/厂家、执行标准、检测资料及任何技术偏离。",
        "duration": "供货周期须满足项目总体进度安排和采购方书面通知要求。报价单位应在报价文件中明确生产、备货、运输及到场周期。",
        "technical": "报价单位应重点核对名称、规格型号、单位、数量、品牌档次、执行标准、是否含运输卸货、检测资料及报价范围边界；清单中列为待补充确认的项目暂不作为正式报价依据。",
    }


def quote_values(row: dict[str, Any], index: int) -> list[Any]:
    return [
        index,
        clean(row.get("来源行号")),
        clean(row.get("追溯号")),
        clean(row.get("标准材料名称")),
        clean(row.get("规格型号")),
        clean(row.get("单位")),
        row.get("合并数量"),
        "",
        "",
        "",
        "",
        "",
        "",
        "",
    ]


def confirm_values(row: dict[str, Any], index: int) -> list[Any]:
    return [
        index,
        clean(row.get("来源行号")),
        clean(row.get("追溯号")),
        clean(row.get("标准材料名称")),
        clean(row.get("规格型号")),
        clean(row.get("单位")),
        row.get("合并数量"),
        clean(row.get("是否需人工确认")),
        clean(row.get("是否可直接发供应商询价")),
        clean(row.get("询价前需补充信息")),
        clean(row.get("备注")),
    ]


def write_docx(path: Path, package: str, rows: list[dict[str, Any]], template: dict[str, str], common: dict[str, str]) -> None:
    direct = ready_rows(rows)
    pending = pending_rows(rows)
    supplier_type = Counter(clean(row.get("供应商类型")) for row in rows).most_common(1)[0][0]
    doc = Document()
    setup_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(template["package_title"])
    style_run(r, bold=True, size=18, font="SimHei")

    add_body(doc, "致：各受邀报价单位")
    add_body(doc, f"我司拟就本项目 {package} 相关内容开展询价工作，请贵司结合本邀请函、报价清单及后续补充技术资料进行报价。")

    add_heading(doc, "一、项目简介")
    add_label(doc, "项目名称：", common.get("project_name", "待补充"))
    add_label(doc, "工程地点：", common.get("project_location", "待补充"))
    add_label(doc, "工程名称/采购包名称：", package)
    add_label(doc, "供应商类型：", supplier_type)
    add_label(doc, "报价范围：", template["scope"])

    add_heading(doc, "二、范围及承包方式")
    add_body(doc, template["contract_mode"])
    add_body(doc, "报价应包含完成本报价范围所需的材料、加工制作、包装、运输、装卸、检测资料、税费、管理费、利润及其他必要费用；如报价单位认为存在遗漏或边界不清，应在报价备注中明确说明。")

    add_heading(doc, "三、质量要求")
    add_body(doc, template["quality"])

    add_heading(doc, "四、工期要求")
    add_body(doc, template["duration"])

    add_heading(doc, "五、报价文件需包含内容")
    for text in [
        "按附件格式填写的报价清单，报价单价、合价、品牌/厂家、税率、供货周期、付款条件应完整。",
        "营业执照、授权或经销证明、类似项目业绩及必要的产品检测报告。",
        "技术响应说明，包含规格、品牌、执行标准、偏离事项及需采购方确认的问题。",
        "联系人、联系电话、报价有效期及售后服务承诺。",
    ]:
        add_body(doc, text)

    add_heading(doc, "六、付款方式")
    add_body(doc, common.get("payment_terms", "待补充"))

    add_heading(doc, "七、联系人")
    add_label(doc, "联系人：", common.get("contact_person", "待补充"))
    add_label(doc, "联系电话：", common.get("contact_phone", "待补充"))
    add_label(doc, "报价截止时间：", common.get("deadline", "待补充"))
    add_label(doc, "报价文件提交方式：", common.get("submit_method", "待补充"))

    add_heading(doc, "八、技术要求 / 询价前说明")
    add_body(doc, template["technical"])
    if pending:
        add_body(doc, f"本包共有 {len(pending)} 项暂不建议直接发供应商正式报价，需先补充或确认关键参数，详见本文件“待补充确认项”。")
    if direct:
        add_body(doc, f"本包共有 {len(direct)} 项可先进入正式报价清单，报价单位应逐项报价并保留追溯号。")
    else:
        add_body(doc, "本包当前暂无可直接发供应商正式报价的清单项，建议先完成参数及技术附件确认后再对外询价。")

    add_heading(doc, "九、附件说明")
    add_body(doc, "附件一：正式报价清单")
    add_body(doc, "附件二：待补充确认项")
    add_body(doc, "附件三：图纸、系统图、技术规格书及其他附件待后续补充确认。")

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run("附件一：正式报价清单")
    style_run(rr, bold=True, size=14, font="SimHei")

    table = doc.add_table(rows=1, cols=len(QUOTE_HEADERS))
    for idx, header in enumerate(QUOTE_HEADERS):
        table.rows[0].cells[idx].text = header
    if direct:
        for idx, row in enumerate(direct, 1):
            cells = table.add_row().cells
            for col, value in enumerate(quote_values(row, idx)):
                cells[col].text = clean(value)
    format_docx_table(table, font_size=7)
    add_body(doc, "说明：报价单价、报价合价、品牌/厂家、税率、供货周期、付款条件及报价备注由报价单位填写；追溯号用于后续报价回填，请勿删除或改写。")

    if pending:
        add_heading(doc, "附件二：待补充确认项")
        table2 = doc.add_table(rows=1, cols=len(CONFIRM_HEADERS))
        for idx, header in enumerate(CONFIRM_HEADERS):
            table2.rows[0].cells[idx].text = header
        for idx, row in enumerate(pending, 1):
            cells = table2.add_row().cells
            for col, value in enumerate(confirm_values(row, idx)):
                cells[col].text = clean(value)
        format_docx_table(table2, header_fill="FCE4D6", font_size=7)

    doc.save(path)


def style_workbook_sheet(ws) -> None:
    dark = PatternFill("solid", fgColor="1F4E78")
    blue = PatternFill("solid", fgColor="5B9BD5")
    orange = PatternFill("solid", fgColor="FCE4D6")
    thin = Side(style="thin", color="D9E2F3")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
    if ws.max_column:
        for cell in ws[1]:
            cell.font = Font(bold=True, size=14, color="FFFFFF")
            cell.fill = dark
            cell.alignment = Alignment(horizontal="center", vertical="center")
        for cell in ws[3]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = blue
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in ws.iter_rows(min_row=4):
        if ws.title in ["待补充确认项", "处理说明"]:
            if ws.title == "待补充确认项":
                for cell in row:
                    cell.fill = orange
    ws.freeze_panes = "A4"
    ws.sheet_view.showGridLines = False
    if ws.max_row >= 3:
        ws.auto_filter.ref = f"A3:{get_column_letter(ws.max_column)}{ws.max_row}"


def write_sheet(wb: Workbook, name: str, title: str, headers: list[str], rows: list[list[Any]], widths: dict[str, int]) -> None:
    ws = wb.create_sheet(name)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    ws.cell(1, 1, title)
    for col, header in enumerate(headers, 1):
        ws.cell(3, col, header)
    for r_idx, row in enumerate(rows, 4):
        for c_idx, value in enumerate(row, 1):
            ws.cell(r_idx, c_idx, value)
    style_workbook_sheet(ws)
    for idx, header in enumerate(headers, 1):
        ws.column_dimensions[get_column_letter(idx)].width = widths.get(header, 16)


def write_xlsx(path: Path, package: str, rows: list[dict[str, Any]], template: dict[str, str]) -> None:
    direct = ready_rows(rows)
    pending = pending_rows(rows)
    wb = Workbook()
    wb.remove(wb.active)

    summary = [
        ["采购包", package],
        ["文件类型", template["package_title"]],
        ["清单总项", len(rows)],
        ["正式报价项", len(direct)],
        ["待补充确认项", len(pending)],
        ["说明", "正式报价项可进入供应商报价清单；待补充确认项建议先补齐参数、技术附件或范围边界后再对外询价。"],
    ]
    write_sheet(wb, "处理说明", "处理说明", ["项目", "内容"], summary, {"项目": 24, "内容": 90})

    quote_rows = [quote_values(row, idx) for idx, row in enumerate(direct, 1)]
    write_sheet(wb, "正式报价清单", "正式报价清单", QUOTE_HEADERS, quote_rows, {
        "序号": 8, "来源行号": 12, "追溯号": 24, "名称": 38, "规格型号": 32, "单位": 10, "数量": 14,
        "品牌/厂家": 16, "税率": 12, "供货周期": 14, "付款条件": 18, "报价单价": 14, "报价合价": 14, "报价备注": 24,
    })

    confirm_rows = [confirm_values(row, idx) for idx, row in enumerate(pending, 1)]
    write_sheet(wb, "待补充确认项", "待补充确认项", CONFIRM_HEADERS, confirm_rows, {
        "序号": 8, "来源行号": 12, "追溯号": 24, "名称": 38, "规格型号": 32, "单位": 10, "数量": 14,
        "是否需人工确认": 16, "是否可直接发供应商询价": 20, "询价前需补充信息": 70, "备注": 42,
    })
    wb.save(path)


def parse_packages(value: str | None, default_packages: list[str]) -> list[str]:
    if not value:
        return default_packages
    if value.strip().lower() in ["all", "全部"]:
        return ["__ALL__"]
    return [part.strip() for part in value.split(",") if part.strip()]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="根据拆分总表生成采购包报价邀请包")
    parser.add_argument("--config", default=str(CONFIG_DIR / "invitation_templates.json"), help="邀请包模板配置")
    parser.add_argument("--standard", default=None, help="采购包拆分总表 xlsx")
    parser.add_argument("--packages", default=None, help="逗号分隔的采购包名称")
    parser.add_argument("--base-dir", default=str(PROJECT_DIR), help="项目根目录")
    parser.add_argument("--output-dir", default=None, help="输出目录")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    base_dir = Path(args.base_dir).resolve()
    cfg = load_json(Path(args.config))
    standard_path = resolve_path(base_dir, args.standard or cfg["default_standard_file"])
    output_root = resolve_path(base_dir, args.output_dir or cfg["default_output_dir"])
    all_rows = read_rows(standard_path)
    packages = parse_packages(args.packages, cfg["default_packages"])
    if packages == ["__ALL__"]:
        packages = [package for package, _ in Counter(clean(row.get("询价采购包")) for row in all_rows).most_common()]

    run_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = output_root / f"{run_id}_采购包邀请清单打样"
    output_dir.mkdir(parents=True, exist_ok=True)
    log: list[str] = [
        f"运行时间：{run_id}",
        f"拆分总表：{standard_path}",
        f"输出目录：{output_dir}",
    ]

    for package in packages:
        rows = [row for row in all_rows if clean(row.get("询价采购包")) == package]
        if not rows:
            log.append(f"{package}：未找到清单项")
            continue
        template = cfg["templates"].get(package) or generic_template(package, rows)
        package_dir = output_dir / package
        package_dir.mkdir(parents=True, exist_ok=True)
        docx_path = package_dir / f"{template['package_title']}.docx"
        xlsx_path = package_dir / f"{template['package_title']}-报价清单.xlsx"
        write_docx(docx_path, package, rows, template, cfg["common"])
        write_xlsx(xlsx_path, package, rows, template)
        log.append(f"{package}：总项{len(rows)}，正式报价项{len(ready_rows(rows))}，待补充确认项{len(pending_rows(rows))}")
        log.append(f"  Word：{docx_path}")
        log.append(f"  Excel：{xlsx_path}")

    (output_dir / "处理日志.txt").write_text("\n".join(log) + "\n", encoding="utf-8")
    print("\n".join(log))


if __name__ == "__main__":
    main()
